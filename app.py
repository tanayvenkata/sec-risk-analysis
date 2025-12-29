"""
Meta 10-K Risk Factor Search - Comparison-Aware Search with LLM Synthesis
"""

import json
import os
import pickle
from dotenv import load_dotenv

load_dotenv()
from pathlib import Path
from collections import Counter

import faiss
import numpy as np
import streamlit as st
from openai import OpenAI
from sentence_transformers import SentenceTransformer
from docx import Document
from docx.shared import Pt
from fpdf import FPDF
import io

# Config
INDEX_DIR = "vector_store"
INDEX_FILE = "faiss_index.bin"
METADATA_FILE = "metadata.pkl"

# OpenRouter config
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")
if not OPENROUTER_API_KEY:
    st.error("OPENROUTER_API_KEY environment variable is required. Set it in your .env file.")
    st.stop()
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"
LLM_MODEL = "google/gemini-2.0-flash-001"  # Fast and cheap for summaries

# Confidence thresholds (for search relevance)
HIGH_CONFIDENCE = 0.75      # > 0.75 = high confidence match
MEDIUM_CONFIDENCE = 0.55    # 0.55-0.75 = medium, < 0.55 = low


# Initialize OpenRouter client
@st.cache_resource
def get_llm_client():
    return OpenAI(api_key=OPENROUTER_API_KEY, base_url=OPENROUTER_BASE_URL)


def generate_search_summary(query: str, results: list, companies: list = None) -> str:
    """Generate a brief summary of search results."""
    if not results:
        return None

    client = get_llm_client()

    # Build dynamic company string
    if companies and len(companies) == 1:
        company_str = companies[0]
        company_possessive = f"{companies[0]}'s"
    elif companies and len(companies) > 1:
        company_str = ", ".join(companies)
        company_possessive = f"these companies' ({company_str})"
    else:
        # Derive from results if not provided
        result_companies = list(set(r.get('company', 'Unknown') for r in results))
        company_str = ", ".join(result_companies) if result_companies else "the selected companies"
        company_possessive = f"{result_companies[0]}'s" if len(result_companies) == 1 else f"these companies' ({company_str})"

    # Format chunks for the prompt (include company in each chunk)
    chunks_text = ""
    for i, r in enumerate(results[:10], 1):  # Limit to top 10 for prompt
        company_tag = r.get('company', '')
        chunks_text += f"[{i}] {company_tag} FY{r['fiscal_year']}: {r['content'][:300]}...\n\n"

    prompt = f"""You are analyzing 10-K Risk Factor disclosures from {company_str} for an equity research analyst.

Query: "{query}"

Here are the top search results:

{chunks_text}

Provide a structured summary:

**Overview** (2-3 sentences): What do {company_possessive} filings disclose about "{query}"? Which fiscal years cover this?

**Key Excerpts:**
- [1]: Brief description of what this excerpt covers
- [2]: Brief description of what this excerpt covers
- [3]: Brief description of what this excerpt covers
(Continue for top 5-6 most relevant)

**Patterns/Trends** (1-2 sentences): Any notable evolution or patterns across years?

Keep the entire response to 1-2 short paragraphs. Be concise but complete."""

    try:
        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000,
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Summary unavailable: {str(e)}"


def generate_comparison_summary(query: str, results_a: list, results_b: list, year_a: int, year_b: int, companies: list = None) -> str:
    """
    Generate a comparison summary using LLM.

    The LLM compares excerpts from both years and identifies:
    - New topics in year_a
    - Removed topics from year_b
    - Wording changes (same topic, different language)
    """
    client = get_llm_client()

    # Build dynamic company string
    if companies and len(companies) == 1:
        company_str = companies[0]
        company_action = companies[0]  # "AAPL added...", "GOOG removed..."
    elif companies and len(companies) > 1:
        company_str = ", ".join(companies)
        company_action = "The companies"
    else:
        # Derive from results if not provided
        result_companies = list(set(r.get('company', 'Unknown') for r in results_a + results_b))
        company_str = ", ".join(result_companies) if result_companies else "the selected companies"
        company_action = result_companies[0] if len(result_companies) == 1 else "The companies"

    # Format excerpts from newer year (include company tag)
    excerpts_a = ""
    for i, r in enumerate(results_a[:8], 1):
        company_tag = r.get('company', '')
        excerpts_a += f"[A{i}] ({company_tag}) {r['content'][:400]}...\n\n"

    # Format excerpts from older year (include company tag)
    excerpts_b = ""
    for i, r in enumerate(results_b[:8], 1):
        company_tag = r.get('company', '')
        excerpts_b += f"[B{i}] ({company_tag}) {r['content'][:400]}...\n\n"

    prompt = f"""You are comparing 10-K Risk Factor disclosures from {company_str} between FY{year_a} and FY{year_b} for an equity research analyst.

Query: "{query}"

**FY{year_a} excerpts ({len(results_a)} found):**
{excerpts_a if excerpts_a else "(No relevant excerpts found)"}

**FY{year_b} excerpts ({len(results_b)} found):**
{excerpts_b if excerpts_b else "(No relevant excerpts found)"}

Analyze these excerpts and identify:

1. **New in FY{year_a}**: Topics or concerns that appear in FY{year_a} but not FY{year_b}
2. **Removed from FY{year_b}**: Topics that were in FY{year_b} but not FY{year_a}
3. **Wording Changes**: Same topic but different language — flag intensification (e.g., "significant" → "intense and increasing"), softening, or shifted emphasis. Quote the specific wording from each year.
4. **Unchanged**: Topics covered similarly in both years (mention briefly)

For each finding, cite the specific excerpt numbers [A1], [B2], etc.

Format your response as:

**What Changed** (lead with this — 2-3 sentences): Start with the most significant difference between FY{year_a} and FY{year_b} regarding "{query}". Be direct: "{company_action} added...", "{company_action} removed...", "{company_action} intensified language about..."

**Key Changes:**
- [NEW] What's in FY{year_a} that wasn't in FY{year_b}... [cite A1]
- [WORDING] "old phrase" → "new phrase" — note if language intensified or softened [cite B1, A3]
- [REMOVED] What was in FY{year_b} but not FY{year_a}... [cite B2]

**Bottom Line** (1 sentence): What should an analyst take away from these changes?

Be specific about wording differences. Quote exact phrases when noting changes."""

    try:
        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Summary unavailable: {str(e)}"


def generate_company_comparison_summary(query: str, results_a: list, results_b: list,
                                        company_a: str, company_b: str, year: int) -> str:
    """Compare what two companies say about a topic in the same fiscal year."""
    client = get_llm_client()

    # Format excerpts from Company A
    excerpts_a = ""
    for i, r in enumerate(results_a[:6], 1):
        excerpts_a += f"[A{i}] {r['content'][:400]}...\n\n"

    # Format excerpts from Company B
    excerpts_b = ""
    for i, r in enumerate(results_b[:6], 1):
        excerpts_b += f"[B{i}] {r['content'][:400]}...\n\n"

    prompt = f"""Compare how {company_a} and {company_b} discuss "{query}" in their FY{year} 10-K Risk Factors.

**{company_a} excerpts ({len(results_a)} found):**
{excerpts_a if excerpts_a else "(No relevant excerpts found)"}

**{company_b} excerpts ({len(results_b)} found):**
{excerpts_b if excerpts_b else "(No relevant excerpts found)"}

Provide a brief comparison:

**Coverage:** Which company dedicates more disclosure to this topic? ({company_a}: {len(results_a)} excerpts, {company_b}: {len(results_b)} excerpts)

**Key Differences:** What does {company_a} mention that {company_b} doesn't, and vice versa? Be specific - cite excerpt numbers [A1], [B2], etc.

**Tone:** Any notable differences in how each company frames this risk? (Note factual differences in language, don't score "aggressiveness")

Keep response to 3-4 sentences. Let the analyst draw conclusions."""

    try:
        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800,
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Summary unavailable: {str(e)}"


@st.cache_resource
def load_model_and_index():
    """Load the embedding model and FAISS index (cached)."""
    with open(f"{INDEX_DIR}/config.json") as f:
        config = json.load(f)

    model = SentenceTransformer(config["model_name"])
    index = faiss.read_index(f"{INDEX_DIR}/{INDEX_FILE}")

    with open(f"{INDEX_DIR}/{METADATA_FILE}", "rb") as f:
        metadata = pickle.load(f)

    return model, index, metadata


def search(query: str, model, index, metadata, top_k: int = 15, year_filter: int = None, company_filter: list = None, section_filter: list = None):
    """Standard search with optional filters."""
    query_embedding = model.encode([query], normalize_embeddings=True)
    distances, indices = index.search(query_embedding.astype(np.float32), k=top_k * 5)

    results = []
    for idx, score in zip(indices[0], distances[0]):
        if idx < 0:
            continue

        meta = metadata[idx]

        # Apply filters
        if year_filter and meta["fiscal_year"] != year_filter:
            continue
        if company_filter and meta["company"] not in company_filter:
            continue
        if section_filter:
            section = meta["section"]
            section_lower = section.lower()
            section_match = False
            for f in section_filter:
                if f in section:
                    section_match = True
                    break
                # "Risk Factors" filter matches any risk-related section (case-insensitive)
                if f == "Risk Factors" and "risk" in section_lower:
                    section_match = True
                    break
                # "MDA" filter matches any MD&A section
                if f == "MDA" and "mda" in section_lower:
                    section_match = True
                    break
            if not section_match:
                continue

        results.append({
            "score": float(score),
            "chunk_id": meta["chunk_id"],
            "company": meta["company"],
            "fiscal_year": meta["fiscal_year"],
            "filed_at": meta.get("filed_at", ""),
            "section": meta["section"],
            "chunk_type": meta["chunk_type"],
            "content": meta["content"],
            "parent_context": meta.get("parent_context", ""),
        })

        if len(results) >= top_k:
            break

    return results


def search_all_years(query: str, model, index, metadata, top_k_per_year: int = 10):
    """Search and return results grouped by year with counts."""
    query_embedding = model.encode([query], normalize_embeddings=True)
    # Get more results to ensure coverage across years
    distances, indices = index.search(query_embedding.astype(np.float32), k=top_k_per_year * 10)

    # Group by year
    results_by_year = {}
    for idx, score in zip(indices[0], distances[0]):
        if idx < 0:
            continue

        meta = metadata[idx]
        year = meta["fiscal_year"]

        if year not in results_by_year:
            results_by_year[year] = []

        if len(results_by_year[year]) < top_k_per_year:
            results_by_year[year].append({
                "score": float(score),
                "chunk_id": meta["chunk_id"],
                "company": meta["company"],
                "fiscal_year": meta["fiscal_year"],
                "section": meta["section"],
                "chunk_type": meta["chunk_type"],
                "content": meta["content"],
                "parent_context": meta.get("parent_context", ""),
            })

    return results_by_year


def get_hit_counts(results: list) -> dict:
    """Get hit counts by year from search results."""
    counts = Counter(r["fiscal_year"] for r in results)
    return dict(sorted(counts.items()))


def export_to_word(query: str, summary: str, results: list, mode: str, metadata: dict = None) -> bytes:
    """Export results to Word document."""
    doc = Document()

    # Title
    doc.add_heading(f"Meta 10-K Risk Factor {mode}", 0)

    # Query
    doc.add_heading("Query", level=1)
    doc.add_paragraph(query)

    # Metadata (hit counts, years compared, etc.)
    if metadata:
        doc.add_heading("Overview", level=1)
        for key, value in metadata.items():
            doc.add_paragraph(f"{key}: {value}")

    # Summary
    if summary:
        doc.add_heading("AI Summary", level=1)
        doc.add_paragraph(summary)

    # Sources
    doc.add_heading("Sources", level=1)
    for i, r in enumerate(results, 1):
        ref = r.get('ref', str(i))
        p = doc.add_paragraph()
        run = p.add_run(f"[{ref}] {r['company']} FY{r['fiscal_year']}")
        run.bold = True
        p.add_run(f" · Score: {r['score']:.2f}")
        doc.add_paragraph(r['section'], style='Caption')
        content = r['content'][:600] + ('...' if len(r['content']) > 600 else '')
        doc.add_paragraph(content)
        doc.add_paragraph()  # spacer

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_pdf(query: str, summary: str, results: list, mode: str, metadata: dict = None) -> bytes:
    """Export results to PDF document."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, f"Meta 10-K Risk Factor {mode}", ln=True)
    pdf.ln(5)

    # Query
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Query:", ln=True)
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, query)
    pdf.ln(3)

    # Metadata
    if metadata:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Overview:", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for key, value in metadata.items():
            pdf.cell(0, 6, f"{key}: {value}", ln=True)
        pdf.ln(3)

    # Summary
    if summary:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "AI Summary:", ln=True)
        pdf.set_font("Helvetica", "", 10)
        # Clean summary for PDF (remove markdown)
        clean_summary = summary.replace('**', '').replace('*', '').replace('#', '')
        pdf.multi_cell(0, 5, clean_summary)
        pdf.ln(3)

    # Sources
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Sources:", ln=True)
    pdf.ln(2)
    for i, r in enumerate(results, 1):
        ref = r.get('ref', str(i))
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 5, f"[{ref}] {r['company']} FY{r['fiscal_year']} - Score: {r['score']:.2f}", ln=True)
        pdf.set_font("Helvetica", "I", 9)
        pdf.cell(0, 5, r['section'][:80], ln=True)
        pdf.set_font("Helvetica", "", 9)
        content = r['content'][:500] + ('...' if len(r['content']) > 500 else '')
        pdf.multi_cell(0, 4, content)
        pdf.ln(2)

    return bytes(pdf.output())


def render_collapsible_summary(summary: str, title: str = "Summary", expanded: bool = False, allow_html: bool = False):
    """Render a collapsible summary with preview."""
    if not summary:
        return

    # Extract first sentence or first 150 chars as preview
    preview = summary.split('\n')[0][:150]
    if len(summary) > 150:
        preview += "..."

    label = f"{title} (Collapse)" if expanded else f"{title} (Expand)"
    with st.expander(label, expanded=expanded):
        st.markdown(summary, unsafe_allow_html=allow_html)

    # Show brief preview outside expander
    if not expanded:
        st.caption(f"*{preview}*")


def get_confidence_meta(score):
    """Return label, css class, and shape class based on score."""
    if score >= 0.75:
        return "High Relevance", "confidence-high", "shape-square"
    elif score >= 0.55:
        return "Medium Relevance", "confidence-med", "shape-triangle"
    else:
        return "Low Relevance", "confidence-low", "shape-circle"


def render_source_chunk(result: dict, ref_num: int, year_label: str = None):
    """Render a source chunk using the accessible Card component."""

    label, confidence_class, shape_class = get_confidence_meta(result['score'])

    year_display = f"FY{result['fiscal_year']}"
    if year_label:
        year_display = f"{year_label} ({year_display})"

    # Format filing date if available
    filed_display = ""
    if result.get('filed_at'):
        try:
            from datetime import datetime
            filed_dt = datetime.fromisoformat(result['filed_at'].replace('Z', '+00:00'))
            filed_display = f"Filed {filed_dt.strftime('%b %Y')}"
        except:
            pass

    # Escape HTML content to prevent rendering issues if content has tags
    import html as html_lib
    safe_content = html_lib.escape(result['content'][:600])
    if len(result['content']) > 600:
        safe_content += "..."

    html = f"""
    <div class="result-card">
        <div class="card-header">
            <div class="card-meta">
                <span class="fiscal-year-badge">{year_display}</span>
                <span class="section-label">{result['section']}</span>
                {f'<span style="color: #888; font-size: 0.75rem;">{filed_display}</span>' if filed_display else ''}
            </div>
            <div class="confidence-indicator {confidence_class} {shape_class}"
                 title="Score: {result['score']:.2f}"
                 aria-label="{label}, Score: {result['score']:.2f}">
                {label}
            </div>
        </div>
        <div class="excerpt-text">
            {safe_content}
        </div>
        <div style="margin-top: 1rem; font-size: 0.8rem; color: #666; display: flex; justify-content: space-between;">
            <span><strong>Ref [{ref_num}]</strong> • ID: {result['chunk_id']}</span>
        </div>
    </div>
    """
    st.markdown(html, unsafe_allow_html=True)


# Page config
st.set_page_config(
    page_title="Meta 10-K Analysis Platform",
    page_icon=None,
    layout="wide"
)

def load_css():
    with open("assets/custom.css") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

load_css()

# Header
st.title("10-K Risk Analysis Platform")

# Load resources first (needed for dynamic header)
with st.spinner("Loading model and index..."):
    model, index, metadata = load_model_and_index()

# Get available years and companies
years = sorted(set(m["fiscal_year"] for m in metadata))
all_companies = sorted(set(m["company"] for m in metadata))

# Format companies for display
if len(all_companies) <= 2:
    companies_display = " & ".join(all_companies)
else:
    companies_display = f"{len(all_companies)} Companies"

# Institutional Dashboard Header (now with dynamic values)
col1, col2, col3 = st.columns(3)
with col1:
    st.markdown("""
    <div class="metric-container">
        <div class="metric-label">Coverage</div>
        <div class="metric-value">Risk Factors + MD&A</div>
        <div style="font-size: 0.7rem; color: #888; margin-top: 4px;">Item 1A & Item 7</div>
    </div>
    """, unsafe_allow_html=True)
with col2:
    year_range = f"{min(years)}-{max(years)}" if years else "N/A"
    st.markdown(f"""
    <div class="metric-container">
        <div class="metric-label">Fiscal Years</div>
        <div class="metric-value">{year_range}</div>
    </div>
    """, unsafe_allow_html=True)
with col3:
    st.markdown(f"""
    <div class="metric-container">
        <div class="metric-label">Companies</div>
        <div class="metric-value">{companies_display}</div>
        <div style="font-size: 0.7rem; color: #888; margin-top: 4px;">{', '.join(all_companies)}</div>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")

# Session state for mode and query
# Modes: "search", "compare_years", "compare_companies"
if "mode" not in st.session_state:
    st.session_state.mode = "search"
if "selected_query" not in st.session_state:
    st.session_state.selected_query = ""

# Sidebar header
st.sidebar.header("10-K Risk Analytics")

# Derive section types from metadata (sections starting with "MDA:" are MD&A, others are Risk Factors)
section_types_found = set()
for m in metadata:
    if m["section"].startswith("MDA:"):
        section_types_found.add("MDA")
    else:
        section_types_found.add("Risk Factors")
all_sections = sorted(section_types_found)

# Company filter with "All Companies" toggle
select_all_companies = st.sidebar.checkbox("All Companies (sector view)", value=False)
if select_all_companies:
    selected_companies = all_companies
    st.sidebar.multiselect(
        "Companies",
        options=all_companies,
        default=all_companies,
        disabled=True
    )
else:
    selected_companies = st.sidebar.multiselect(
        "Companies",
        options=all_companies,
        default=[all_companies[0]] if all_companies else []
    )

# Section filter
selected_sections = st.sidebar.multiselect(
    "Sections",
    options=all_sections,
    default=["Risk Factors"]
)

st.sidebar.markdown("---")

# Unified Filter Toolbar
year_options = [f"FY{y}" for y in reversed(years)]

# Mode selector
mode_col1, mode_col2, mode_col3, mode_col4 = st.columns([1.5, 1.5, 2, 3])
with mode_col1:
    if st.button("Search", type="primary" if st.session_state.mode == "search" else "secondary", use_container_width=True):
        st.session_state.mode = "search"
        st.rerun()
with mode_col2:
    if st.button("Compare Years", type="primary" if st.session_state.mode == "compare_years" else "secondary", use_container_width=True):
        st.session_state.mode = "compare_years"
        st.rerun()
with mode_col3:
    if st.button("Compare Companies", type="primary" if st.session_state.mode == "compare_companies" else "secondary", use_container_width=True):
        st.session_state.mode = "compare_companies"
        st.rerun()

st.markdown("")  # Spacer

# Mode-specific controls
if st.session_state.mode == "compare_years":
    # Compare Years mode: two year selectors
    col1, col2, col3 = st.columns([2, 2, 4])
    with col1:
        year_a = st.selectbox("Newer Year", year_options, index=0)
    with col2:
        year_b = st.selectbox("Older Year", year_options, index=1)
    with col3:
        top_k = st.slider("Result Limit", 5, 20, 10)

    year_a_int = int(year_a.replace("FY", ""))
    year_b_int = int(year_b.replace("FY", ""))
    year_filter = None

elif st.session_state.mode == "compare_companies":
    # Compare Companies mode: two company selectors + year
    col1, col2, col3, col4 = st.columns([2, 2, 2, 2])
    with col1:
        company_a = st.selectbox("Company A", all_companies, index=0)
    with col2:
        # Default to second company if available
        default_b_idx = 1 if len(all_companies) > 1 else 0
        company_b = st.selectbox("Company B", all_companies, index=default_b_idx)
    with col3:
        compare_year = st.selectbox("Fiscal Year", year_options, index=0)
    with col4:
        top_k = st.slider("Results/Company", 5, 15, 8)

    compare_year_int = int(compare_year.replace("FY", ""))
    year_filter = None

else:
    # Search mode: single year filter
    col1, col2 = st.columns([2, 6])
    with col1:
        all_year_options = ["All Years"] + year_options
        selected_year = st.selectbox("Fiscal Year Scope", all_year_options)
        year_filter = None
        if selected_year != "All Years":
            year_filter = int(selected_year.replace("FY", ""))
    with col2:
        top_k = st.slider("Result Limit", 5, 30, 15)

# Query input
if st.session_state.mode == "compare_years":
    query_label = "Compare Topic"
    query_placeholder = "e.g., AI regulation, competition, antitrust..."
elif st.session_state.mode == "compare_companies":
    query_label = "Topic to Compare"
    query_placeholder = "e.g., AI regulation, China risks, antitrust..."
else:
    query_label = "Risk Query"
    query_placeholder = "e.g., AI regulation, Apple competition, China risks..."
query = st.text_input(query_label, value=st.session_state.selected_query, placeholder=query_placeholder)
# Clear selected_query after it's been used
if st.session_state.selected_query:
    st.session_state.selected_query = ""

# Results section - adapts based on mode
if st.session_state.mode == "compare_years":
    # --- COMPARE MODE ---
    if year_a_int == year_b_int:
        st.warning("Please select two different years to compare.")
    elif query:
        st.markdown(f"### {query}: {year_a} vs {year_b}")

        # Search both years
        with st.spinner("Searching both years..."):
            results_a = search(query, model, index, metadata, top_k=top_k, year_filter=year_a_int, company_filter=selected_companies, section_filter=selected_sections)
            results_b = search(query, model, index, metadata, top_k=top_k, year_filter=year_b_int, company_filter=selected_companies, section_filter=selected_sections)

        # Show hit counts
        col1, col2 = st.columns(2)
        with col1:
            st.metric(f"{year_a}", f"{len(results_a)} results")
        with col2:
            st.metric(f"{year_b}", f"{len(results_b)} results")

        # Trend indicator
        if len(results_a) > len(results_b):
            st.success(f"Increased Focus in {year_a} (Coverage expanded)")
        elif len(results_b) > len(results_a):
            st.warning(f"Decreased Focus in {year_a} (Coverage reduced)")
        else:
            st.info("Stable Coverage")

        # Generate comparison summary
        if results_a or results_b:
            with st.spinner("Analyzing differences..."):
                summary = generate_comparison_summary(query, results_a, results_b, year_a_int, year_b_int, companies=selected_companies)

                # Post-process summary for accessible styling
                if summary:
                    summary = summary.replace("[NEW]", '')
                    summary = summary.replace("[REMOVED]", '')
                    summary = summary.replace("[IMPORTANT]", '<span class="diff-added">')
                    summary = summary.replace("[REMOVED]", '<span class="diff-removed">')
                    summary = summary.replace("[/IMPORTANT]", '</span>')
                    summary = summary.replace("[/REMOVED]", '</span>')

            st.markdown("---")
            st.markdown("### Comparison Matrix")
            render_collapsible_summary(summary, title="Change Analysis", expanded=True, allow_html=True)

            st.markdown("---")
            st.markdown("### Sources")

            if results_a:
                st.markdown(f"**{year_a} (Newer)**")
                for i, r in enumerate(results_a, 1):
                    render_source_chunk(r, f"A{i}", "Newer")

            if results_b:
                st.markdown(f"**{year_b} (Older)**")
                for i, r in enumerate(results_b, 1):
                    render_source_chunk(r, f"B{i}", "Older")

            # Export buttons
            st.markdown("---")
            st.markdown("### Export")
            exp_col1, exp_col2 = st.columns(2)

            all_results = [{"ref": f"A{i+1}", **r} for i, r in enumerate(results_a)] + \
                          [{"ref": f"B{i+1}", **r} for i, r in enumerate(results_b)]

            compare_metadata = {
                "Comparison": f"{year_a} vs {year_b}",
                f"{year_a} Results": len(results_a),
                f"{year_b} Results": len(results_b),
            }

            with exp_col1:
                word_bytes = export_to_word(query, summary, all_results, "Comparison", compare_metadata)
                st.download_button(
                    "Download Word",
                    data=word_bytes,
                    file_name=f"meta_compare_{query[:20].replace(' ', '_')}_{year_a}_vs_{year_b}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            with exp_col2:
                pdf_bytes = export_to_pdf(query, summary, all_results, "Comparison", compare_metadata)
                st.download_button(
                    "Download PDF",
                    data=pdf_bytes,
                    file_name=f"meta_compare_{query[:20].replace(' ', '_')}_{year_a}_vs_{year_b}.pdf",
                    mime="application/pdf"
                )
        else:
            st.warning("No results found. Try a different query.")
    else:
        st.info("Enter a topic above to see what changed between the selected years.")

elif st.session_state.mode == "compare_companies":
    # --- COMPARE COMPANIES MODE ---
    if company_a == company_b:
        st.warning("Please select two different companies to compare.")
    elif query:
        st.markdown(f"### {query}: {company_a} vs {company_b} ({compare_year})")

        # Search each company separately
        with st.spinner("Searching both companies..."):
            results_a = search(query, model, index, metadata, top_k=top_k,
                             year_filter=compare_year_int,
                             company_filter=[company_a],
                             section_filter=selected_sections)
            results_b = search(query, model, index, metadata, top_k=top_k,
                             year_filter=compare_year_int,
                             company_filter=[company_b],
                             section_filter=selected_sections)

        # Show hit counts
        col1, col2 = st.columns(2)
        with col1:
            st.metric(company_a, f"{len(results_a)} results")
        with col2:
            st.metric(company_b, f"{len(results_b)} results")

        # Generate comparison summary
        if results_a or results_b:
            with st.spinner("Comparing companies..."):
                summary = generate_company_comparison_summary(
                    query, results_a, results_b, company_a, company_b, compare_year_int
                )

            st.markdown("---")
            render_collapsible_summary(summary, title="Comparison Summary", expanded=True)

            # Two-column results
            st.markdown("---")
            st.markdown("### Sources")

            col_a, col_b = st.columns(2)

            with col_a:
                st.markdown(f"**{company_a}**")
                if results_a:
                    for i, r in enumerate(results_a, 1):
                        render_source_chunk(r, f"A{i}")
                else:
                    st.caption("No results found")

            with col_b:
                st.markdown(f"**{company_b}**")
                if results_b:
                    for i, r in enumerate(results_b, 1):
                        render_source_chunk(r, f"B{i}")
                else:
                    st.caption("No results found")
        else:
            st.warning("No results found for either company. Try a different query.")
    else:
        st.info("Enter a topic above to compare how the two companies discuss it.")

else:
    # --- SEARCH MODE ---
    if query:
        results = search(query, model, index, metadata, top_k=top_k, year_filter=year_filter, company_filter=selected_companies, section_filter=selected_sections)

        if results:
            hit_counts = get_hit_counts(results)

            st.markdown(f"### Found {len(results)} results")

            if len(hit_counts) > 1:
                st.markdown("**FY Trend Analysis:**")
                counts_str = " → ".join([f"FY{y}: **{c}**" for y, c in sorted(hit_counts.items())])
                st.markdown(counts_str)

            with st.spinner("Synthesizing analysis..."):
                summary = generate_search_summary(query, results, companies=selected_companies)
            if summary:
                render_collapsible_summary(summary, title="Executive Synthesis", expanded=True)

            st.markdown("---")
            st.markdown("### Sources")
            st.caption("**Score guide:** 0.7+ = strong match, 0.5-0.7 = moderate, <0.5 = weak")

            for i, r in enumerate(results, 1):
                render_source_chunk(r, i)

            # Export buttons
            st.markdown("---")
            st.markdown("### Export")
            exp_col1, exp_col2 = st.columns(2)

            export_metadata = {
                "Total Results": len(results),
                "Years": ", ".join([f"FY{y}: {c}" for y, c in hit_counts.items()])
            }

            with exp_col1:
                word_bytes = export_to_word(query, summary, results, "Search", export_metadata)
                st.download_button(
                    "Download Word",
                    data=word_bytes,
                    file_name=f"meta_search_{query[:20].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            with exp_col2:
                pdf_bytes = export_to_pdf(query, summary, results, "Search", export_metadata)
                st.download_button(
                    "Download PDF",
                    data=pdf_bytes,
                    file_name=f"meta_search_{query[:20].replace(' ', '_')}.pdf",
                    mime="application/pdf"
                )
        else:
            st.warning("No results found. Try a different query or remove filters.")

# Sidebar footer
with st.sidebar:
    st.markdown("---")
    st.markdown("### Trending Topics")
    st.caption("Click to search")

    if st.session_state.mode in ["compare_years", "compare_companies"]:
        topics = ["AI regulation", "competition", "antitrust", "workforce reduction", "China"]
    else:
        topics = ["AI regulation", "Apple iOS ATT", "China risks", "FTC antitrust", "workforce reduction", "metaverse"]

    for topic in topics:
        if st.button(topic, key=f"topic_{topic}", use_container_width=True):
            st.session_state.selected_query = topic
            st.rerun()

    st.markdown("---")
    st.markdown("### System Methodology")
    if st.session_state.mode == "compare_years":
        st.markdown("""
        1. Search both years for your query
        2. LLM compares excerpts directly
        3. Identifies new, removed, and wording changes
        4. All sources cited for verification
        """)
    elif st.session_state.mode == "compare_companies":
        st.markdown("""
        1. Search same topic in both companies
        2. Display results side-by-side
        3. LLM summarizes coverage differences
        4. All sources cited for verification
        """)
    else:
        st.markdown("""
        1. Search finds relevant excerpts
        2. Shows hit counts by year (trend signal)
        3. LLM synthesizes findings
        4. Sources provided for verification
        """)

    st.markdown("---")
    st.markdown("### Index Stats")
    st.markdown(f"- **Documents:** {index.ntotal:,}")
    st.markdown(f"- **Years:** FY{min(years)}-FY{max(years)}")

    # Scope disclaimer
    st.markdown("---")
    st.caption("""
    **Scope:** Item 1A (Risk Factors) and Item 7 (MD&A) only.
    Does not include other 10-K sections, 10-Q filings, or 8-Ks.
    Always verify against original SEC filings.
    """)
